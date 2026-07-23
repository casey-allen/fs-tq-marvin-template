#!/usr/bin/env python3
"""SkyVault VDR Executive Summary (Word) — python-docx fallback replicating the vdr-review
fixed 5-section structure, black-and-white formatting."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/casey/git/obsidian/Projects/SkyVault-VDR/SkyVault-Executive-Summary.docx"
BLACK = RGBColor(0,0,0); WHITE = RGBColor(0xFF,0xFF,0xFF)
DARK = RGBColor(0x1A,0x1A,0x1A); GRAY = RGBColor(0x59,0x59,0x59)

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = BLACK
for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def setcell(cell, text, bold=False, color=BLACK, size=9, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name='Calibri'
    if fill: shade(cell, fill)
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)

def heading(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'2'); bottom.set(qn('w:color'),'595959')
    pbdr.append(bottom); pPr.append(pbdr)

def verdict(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BLACK
    return p

def body(text, italic=False, color=BLACK):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(10); r.font.color.rgb = color
    return p

def mktable(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i,h in enumerate(headers): setcell(t.rows[0].cells[i], h, bold=True, color=WHITE, size=9, fill="1A1A1A")
    for ri,row in enumerate(rows):
        cells = t.add_row().cells
        fillc = "F8F8F8" if ri%2 else "FFFFFF"
        for i,val in enumerate(row): setcell(cells[i], val, size=9, fill=fillc)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width = Inches(w)
    return t

# ── Title ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("FLUIDSTACK — VDR EXECUTIVE SUMMARY"); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=DARK
p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Project SkyVault — Temple Strategic Power Offtake (BKV / Banpu)"); r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=BLACK

# ── 1. Deal Snapshot ──
heading("1. Deal Snapshot")
mktable(["Field","Detail"], [
    ["Project / Codename","Temple Strategic Power Offtake / SkyVault"],
    ["Location","Temple, Bell County, TX"],
    ["ISO / Utility","ERCOT / Oncor (POI: Knob Creek 345 kV)"],
    ["Gross MW","Up to 1,500 MW (750 existing CCGT + 750 newbuild CCGT); +~100-400 MW modular bridge"],
    ["Current Stage","Diligence in progress — competitive Barclays auction (indicative proposals due 12/5/2025)"],
    ["Seller","BKV Corp (NYSE: BKV) + Banpu Power U.S. JV"],
    ["Transaction Structure","Behind-the-meter co-location power offtake (RPSA + Private Use Network) + fee-simple DC land purchase"],
], widths=[1.6,5.4])

# ── 2. Power Verdict ──
heading("2. Power Verdict")
verdict("Power is real but unsecured: the underlying generation is operating (Temple I/II SGIA, 2012) and credible, but there is no executed offtake, no signed data-center load interconnection, no completed ERCOT/Oncor study, no defined CIAC, and the Private Use Network is still conceptual.")
mktable(["Milestone","Status","Date"], [
    ["Interconnection Agreement","Generation: Executed (2012). Load: Interim FEA only (non-binding)","Interim FEA executed Nov 2025; Final FEA — Not in VDR"],
    ["ESA / PPA","Not Initiated — form PSA + non-binding term sheets, blank pricing","N/A"],
    ["First Power","Realistic (conservative): existing units via PUN ~2H 2029 if Oncor PUN ISD (5/25/2029) holds","~Q4 2029, up to 750 MW"],
    ["Full Capacity","Temple III newbuild — Guaranteed Substantial Completion","Aug 2, 2030, +~750 MW"],
    ["CIAC","Unknown — $6.5M interim Oncor security is an explicit placeholder; network upgrades unquantified","N/A"],
], widths=[1.7,3.5,1.8])

# ── 3. Site Verdict ──
heading("3. Site Verdict")
verdict("Site is workable but encumbered: ample land and excellent fiber, but Knob Creek floodplain and >10% slope cut into buildable area, the PUN gen-tie easements are unsecured, environmental RECs (PFAS suspected) are uncharacterized, and no geotech exists.")
mktable(["Dimension","Rating","One-line basis"], [
    ["Buildable Acreage","Claimed","~892 buildable ac (Bowie) gross; net usable cut by floodplain + slope, not quantified"],
    ["Flood","Amber","Knob Creek floodway + Zone AE bisects parcels; footprints sited outside, no CLOMR"],
    ["Environmental","Amber","Phase I done; RECs: arsenic GW, 1995 AFFF/PFAS, abandoned UST, lead. Phase II not done"],
    ["Title","Amber","Per-parcel commitments issued; open Schedule C items; Schwake grantor-capacity issue; CLMG $435M plant mortgage"],
    ["Water","Constrained","City of Temple reclaimed effluent; rate sheet only, no will-serve, no volumes"],
], widths=[1.5,1.0,4.5])

# ── 4. Top Risks ──
heading("4. Top Risks")
mktable(["Risk","Severity","Finding","Mitigation"], [
    ["No executed offtake or load IA","Critical","No RPSA (form + blank-price term sheets); load interconnection at interim, non-binding Oncor FEA; LLIS results not yet received","Open"],
    ["CIAC / upgrades undefined","Critical","$6.5M interim is a placeholder; EPE study shows 750 MW load needs ~13 transmission limiters resolved; total cost unquantified","Open"],
    ["Firm power hinges on CCGTs + gas","High","Grid cannot backstop 750 MW (N-1 import ~10 MW); firm gas (ETC 75,000 MMBtu/d) expires Q4 2027 then month-to-month","Open"],
    ["PFAS / Phase II not done","High","1995 AFFF spill 100 ft up-gradient of Bowie parcel; Phase II recommended, not performed; no sampling data — cannot be ruled in/out","Open"],
    ["Non-exclusive auction + SB6","High","Competitive Barclays process, seller can walk anytime; Texas SB6 large-load approval is a CP and contemplates curtailment; rules still being written","Open"],
], widths=[1.5,0.75,3.75,0.9])

# ── 5. Recommendation ──
heading("5. Recommendation")
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
shadep = p._p.get_or_add_pPr(); sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'1A1A1A'); shadep.append(sh)
r = p.add_run("CONDITIONAL ADVANCE  (Score 4.0 / 10 — no red line tripped)"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=WHITE
body("SkyVault is a high-quality, operating gas asset from a credible public seller (BKV/Banpu) with attractive indicative firm-power pricing ($70-95/MWh) and an unusually deep, well-organized data room — but it is offered through a competitive auction at the indicative-proposal stage, and every power-securing milestone remains open. There is no executed offtake, no signed data-center load interconnection, no completed ERCOT/Oncor study, no defined CIAC, a conceptual PUN, and an SB6 large-load regime that explicitly contemplates curtailment. The gaps are early-stage rather than fatal, which justifies advancing to compete — but only conditionally, and with eyes open to a power thesis that depends on the on-site CCGTs running (the grid cannot firm 750 MW behind them) and on a 2029-2030 interconnection path that no utility has yet committed.")
body("Advance to participate in the process, conditioned on the seller delivering — pre-LOI/binding-bid — the following objectively verifiable items:")
for i,c in enumerate([
    "Executed (or substantially negotiated) RPSA with filled-in energy/capacity pricing, a committed availability guarantee, and an LD schedule.",
    "Completed ERCOT/Oncor load study (LLIS results, SIS, Final FEA) with a defined network-upgrade scope, CIAC dollar figure, and payment schedule.",
    "Shared Facilities Agreement (co-ownership %, cost allocation) and evidence of ERCOT PUN registration under Nodal Protocol 10.3.2.3.",
    "Phase II ESA with PFAS groundwater sampling results and an NFA letter (or a defined, bounded remediation cost) for the Bowie parcel.",
    "Firm gas transport covering the full combined load through the full offtake term, plus executed PUN gen-tie easements and BNSF/Oncor/TxDOT crossing agreements.",
], 1):
    pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Inches(0.25); pp.paragraph_format.space_after = Pt(2)
    rr = pp.add_run(f"{i}.  {c}"); rr.font.size = Pt(9.5)

# Footer
pf = doc.add_paragraph(); pf.paragraph_format.space_before = Pt(10)
rf = pf.add_run("Based on VDR as of June 15, 2026. See companion Excel workbook (SkyVault-VDR-Tables.xlsx) for the completed RFI, 82-item Risk Matrix, Q&A Log, Deal Record, and Scoring Matrix.")
rf.italic = True; rf.font.size = Pt(8); rf.font.color.rgb = GRAY

doc.save(OUT)
print("Saved:", OUT)
