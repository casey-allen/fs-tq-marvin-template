#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)
HDR_FILL = "1F3864"
ALT_FILL = "EAF0F7"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[h]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = NAVY

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, white=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def make_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, white=True)
        shade(t.rows[0].cells[i], HDR_FILL)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
            if r_idx % 2 == 1:
                shade(cells[i], ALT_FILL)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    return t

def para(text, size=10.5, italic=False, color=None, after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.size = Pt(10.5)
        rest = p.add_run(text)
        rest.font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p

# ---- Title block ----
title = doc.add_paragraph()
tr = title.add_run("APS Large Load Interconnection Process")
tr.font.size = Pt(20); tr.font.bold = True; tr.font.color.rgb = NAVY
sub = doc.add_paragraph()
sr = sub.add_run("Arizona Public Service — Data Center / Extra-Large Load Reference")
sr.font.size = Pt(12); sr.font.color.rgb = GREY
meta = doc.add_paragraph()
mr = meta.add_run("Prepared for Site Selection  ·  Fluidstack  ·  June 2026")
mr.font.size = Pt(9.5); mr.italic = True; mr.font.color.rgb = GREY

# rule
pr = doc.add_paragraph()
pPr = pr._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1F3864")
pbdr.append(bottom); pPr.append(pbdr)

# Confidence banner
para("Confidence basis: Rate terms are GROUNDED in the current APS XHLF tariff (A.C.C. No. 6067, Rev 3, "
     "eff. March 8, 2024, Decision No. 79293). Process stage map and timelines are ESTIMATED — no APS "
     "large-load process guide or facilities study is on file. Collateral and minimum-bill amounts are set per "
     "Electric Supply Agreement and are not publicly fixed. Verify against live ACC dockets before any commitment.",
     size=9.5, italic=True, color=GREY, after=10)

# ---- Section: Context ----
doc.add_heading("1.  Market Context", level=1)
para("Arizona is NOT an organized ISO/RTO market. There is no independent interconnection queue, no LMP energy "
     "market, and no FERC ISO tariff governing load interconnection. APS is vertically integrated and is its own "
     "Balancing Authority within WECC — it owns both the wires and the generation. Practical implications:")
bullet("you negotiate directly with one entity, and “can they serve it” is a generation-adequacy question, not just a wires question.")
bullet("the process is governed by state-filed tariffs and line-extension schedules, regulated by the Arizona Corporation Commission (ACC) — not FERC ISO rules.")
bullet("the ACC can change rates and tariff terms mid-project — a live regulatory risk unique to this structure.")
bullet("Arizona is capacity- and transmission-constrained in high-growth corridors (Phoenix metro, Buckeye, Coolidge, Casa Grande). APS large-customer peak demand ≈ 13.1 GW (2026); pending data-center interest reported at ~10–18 GW.")
bullet("“bring your own power” (customer-assigned new generation) is an explicit policy direction, not a fringe option.")

# ---- Section: Process Map ----
doc.add_heading("2.  Process Map", level=1)
para("Position a project against these stages. The firm gate is Stage 6 — nothing before it is a committed power position.", after=8)
make_table(
    ["#", "Stage", "Typical Duration", "Owner"],
    [
        ["1", "Load request to APS large-load / new-service team (MW, location, ramp, in-service date)", "Customer action", "Customer"],
        ["2", "Preliminary capacity / feasibility screening (can APS serve it at that location/size/timing)", "1–3 months [EST]", "APS"],
        ["3", "Study agreement + deposit → detailed facilities / system-impact study", "3–6 months [EST]", "APS"],
        ["4", "Commercial commitments: ESA negotiation, site control, collateral, ramp + minimum-bill terms", "Costs begin accruing", "Customer / APS"],
        ["5", "ESA + Line Extension Agreement (Service Schedule 3) executed; ACC approval if special contract", "2–6 months [EST]", "APS / ACC"],
        ["6", "CIAC paid + collateral posted → APS orders long-lead equipment  ★ FIRM GATE", "Customer action", "Customer"],
        ["7", "Engineering + construction (substation / transmission)", "24–48+ months [EST]", "APS"],
        ["8", "Energization", "—", "APS"],
    ],
    [0.35, 5.1, 1.45, 0.95],
)
para("", after=2)
para("Total duration: ~36–60 months [ESTIMATED] — gated by transformer lead times and area capacity, not engineering. "
     "Clock starts at the study agreement (Stage 3); hard cost accrual begins at Stage 4.", size=9.5, color=GREY)

# ---- Section: Firm Power ----
doc.add_heading("3.  Path to Firm Power", level=1)
bullet("executed ESA + executed Line Extension Agreement (Service Schedule 3),", bold_lead="Firm gate = ")
bullet("plus CIAC paid / collateral posted,")
bullet("plus ACC approval if the deal is a special contract — which together trigger equipment procurement.")
para("Anything earlier — a feasibility study, a draft ESA — is non-firm and not something APS will build against. "
     "The cheap optionality sits in the study stage; real capital exposure is the CIAC + collateral at the firm gate.", after=8)

# ---- Section: Cost & XHLF terms ----
doc.add_heading("4.  Cost & Rate Terms (XHLF Tariff)", level=1)
para("Grounded in APS Rate Schedule XHLF (A.C.C. No. 6067, Rev 3). XHLF is the large-load rate vehicle for "
     "extra-high-load-factor customers such as data centers.", after=8)

doc.add_heading("Eligibility", level=3)
make_table(
    ["Requirement", "Threshold", "Source"],
    [
        ["Monthly maximum demand", "≥ 5,000 kW", "GROUNDED"],
        ["Load factor", "≥ 92% for ≥ 9 of prior 12 months", "GROUNDED"],
        ["Econ-Dev & Sustainability features", "≥ 15,000 kW + 50% carbon-free activity", "GROUNDED"],
    ],
    [3.3, 3.05, 1.5],
)
para("", after=4)
doc.add_heading("Charges", level=3)
make_table(
    ["Charge", "Transmission", "Primary", "Secondary"],
    [
        ["Energy ($/kWh)", "0.04210", "0.04210", "0.04210"],
        ["Demand ($/kW)", "13.274", "17.290", "18.750"],
        ["Basic service ($/day)", "43.130", "8.833", "5.682"],
    ],
    [2.85, 1.7, 1.65, 1.65],
)
para("", after=2)
bullet("up to 35% off the Generation–Capacity charge, ≤ 6 years, new load only (Service Schedule 9).", bold_lead="Econ-Dev discount: ")
bullet("ICE Palo Verde hub day-ahead on/off-peak + $0.00314/kWh non-bypassable fee (in lieu of standard fuel adjustment).", bold_lead="Optional market-proxy energy: ")
bullet("Econ-Dev/Sustainability program limited to +50 MW/yr per customer, 500 MW total (APS discretion on market liquidity).", bold_lead="Program cap: ")

# ---- Section: Interconnection Process Costs ----
doc.add_heading("5.  Interconnection Process Costs", level=1)
para("Read this first: APS does NOT publish a flat large-LOAD study fee. Load connections run through "
     "Service Schedule 3 (line extensions), where the study charge equals APS’s actual estimated cost to prepare "
     "the work — not a fixed tariff number. The tiered deposits you will find online ($20K / $40K feasibility & "
     "system-impact, or $5K + $100K + $200K) are GENERATION interconnection (Schedule 6 / FERC queue reform) and do "
     "NOT apply to a data-center load. Both are shown below so the distinction is unambiguous.", after=8, bold=False)

doc.add_heading("A.  Large-Load Study & Connection Costs (Service Schedule 3 — applies to data centers)", level=3)
make_table(
    ["Cost item", "Amount / basis", "Notes", "Confidence"],
    [
        ["Preliminary sketch + rough cost estimate", "No charge", "Provided by APS on request (§27.1)", "GROUNDED"],
        ["Special study / detailed design / cost estimate", "= APS’s estimated cost to prepare it (no fixed fee)", "Credited to construction cost if you proceed; NON-REFUNDABLE if you don’t (§27.1)", "GROUNDED"],
        ["Cost-sharing test", "“Economic Feasibility”: projected annual delivery revenue vs. cost", "Determines free allowance vs. refundable advance", "GROUNDED"],
        ["Construction cost above allowance", "Paid in full by applicant", "Due at Line Extension Agreement signing, before APS installs facilities", "GROUNDED"],
        ["CIAC (≥ 15 MW, transmission service)", "Full cost of transmission facilities, in lieu of buying them", "APS may finance at its WACC up to 10 years (XHLF option)", "GROUNDED"],
        ["Collateral & minimum bill", "Set per Electric Supply Agreement — not publicly fixed", "New / higher terms proposed in 2025 rate case (see §6)", "UNKNOWN (per ESA)"],
        ["Long-lead equipment", "Transmission transformers, breakers — 2–4 yr lead", "Drives schedule, not a fee", "MARKET NORM"],
    ],
    [2.15, 2.35, 3.0, 1.2],
)
para("Industry comparable for a load study/evaluation fee (NON-APS, order-of-magnitude only): Tri-State proposes "
     "~$80K @ 45 MW / $150K @ 100 MW / $250K @ 200 MW.", size=9, italic=True, color=GREY, after=8)

doc.add_heading("B.  Generation Interconnection Deposits (Schedule 6 / LGIP — REFERENCE ONLY, not load)", level=3)
para("Listed only so these are not mistaken for load costs. These apply to generators (solar, storage, gas), not to a data-center load.", size=9, italic=True, color=GREY, after=4)
make_table(
    ["Stage / item", "Deposit", "Confidence"],
    [
        ["Feasibility Study (≥ 69 kV)", "$20,000", "GROUNDED"],
        ["System Impact Study (≥ 69 kV)", "$40,000", "GROUNDED"],
        ["Facilities Study", "Per Schedule 6 (actual cost; prior study deposits credited)", "GROUNDED"],
        ["Queue-reform application package", "$5,000 non-refundable app fee + ~$100,000 study deposit + $200,000 commercial-readiness deposit", "GROUNDED"],
        ["Large Generator Interconnection Agreement (LGIA)", "$160,000 (≤ 75 MW) / $250,000 (≥ 76 MW)", "GROUNDED"],
    ],
    [2.6, 4.45, 1.65],
)
para("", after=2)
para("Bottom line: for a data-center load, budget the study as APS’s actual cost to prepare it (credited toward "
     "construction if you build) — the material money is CIAC + collateral at the firm gate, which can run tens of "
     "millions for a large campus, not the study fee.", size=10, bold=True, after=8)

# ---- Section: Regulatory risk ----
doc.add_heading("6.  Regulatory Risk (Active Dockets)", level=1)
bullet("APS proposes >45% rate increase for extra-large users, new minimum-bill and collateral requirements, "
       "“bring your own power” cost assignment, accelerated service via long-term contracts with up-front "
       "infrastructure contributions, and formula rates. PROPOSED, not yet effective.",
       bold_lead="ACC docket E-01345A-25-0105 (2025 rate case): ")
bullet("Commission-wide review of data-center / large-load rate classifications and ratepayer protection. "
       "APS and SRP have proposed their own queue processes and standard agreements requiring up-front financial "
       "support. Terms not yet quantified.",
       bold_lead="ACC docket E-00000A-25-0069 (data-center inquiry): ")
para("Net: cost-allocation, collateral, and minimum-bill terms for Arizona large load are actively in flux through 2026. "
     "Verify against the live docket before committing.", after=8, italic=True, color=GREY)

# ---- Section: Deal notes ----
doc.add_heading("7.  Deal Notes", level=1)
bullet("Generation adequacy is the dominant gating risk, not wires availability — confirm APS can supply energy at the site/size/timing before valuing the wires path.")
bullet("Expect to be pushed toward BYOP / demand-flexibility and a 50% carbon-free commitment to access the Econ-Dev & Sustainability features.")
bullet("The 92% load-factor floor is data-center-friendly (steady load); the tariff exempts billing cycles where a miss is solely due to in-cycle load growth — watch ramp mechanics.")
bullet("Real capital exposure is CIAC + collateral at the firm gate, which can run tens of millions for a large campus.")

# ---- Section: To upgrade ----
doc.add_heading("8.  To Upgrade This Analysis", level=1)
para("Obtain these to move from ESTIMATED to GROUNDED:", after=4)
bullet("APS large-load process guide / a real facilities-study scoping — flips the stage map, study cost, and firm-gate timing. (Highest value.)")
bullet("ACC eDocket E-01345A-25-0105 filings — quantifies proposed collateral, minimum bill, BYOP terms.")
bullet("ACC eDocket E-00000A-25-0069 filings — quantifies data-center-specific rules.")
bullet("SRP large-load process guide + rate schedule, and TEP large-load tariff — for full Arizona coverage.")

# ---- Sources ----
doc.add_heading("Sources", level=2)
for s in [
    "APS Rate Schedule XHLF, A.C.C. No. 6067, Rev 3 (Decision No. 79293), effective March 8, 2024.",
    "APS Service Schedule 3 (Extensions of Electric Distribution Lines and Services), A.C.C. No. 6217, Rev 16, effective March 8, 2024 — §27.1 study/design, Economic Feasibility, allowances.",
    "APS Service Schedule 6 (Interconnection Study, Non-FERC Generation) and APS LGIP/queue-reform filings — source of the generation deposit figures.",
    "APS — “Data Centers: How We’re Protecting Customers While Planning for Big Energy Needs.”",
    "Arizona Corporation Commission — Data Center/Large Load Workshop Highlights (April 2026).",
    "APS 2025 Rate Case (ACC docket E-01345A-25-0105).",
    "ACC data-center inquiry (docket E-00000A-25-0069).",
    "EUCI — Utilities adopt large load tariffs (Tri-State evaluation-fee comparable).",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(s); r.font.size = Pt(9)

out = "/Users/casey/fs-tq-marvin-template/APS_Large_Load_Interconnection_Process.docx"
doc.save(out)
print("Saved:", out)
